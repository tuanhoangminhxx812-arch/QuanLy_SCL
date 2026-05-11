import streamlit as st
import pandas as pd
import os
import re
import datetime
from io import BytesIO
from data_helpers import load_tonghop, load_pm092, get_trang_thai_list
from form_module import load_db_data

st.set_page_config(page_title="Quản lý SCL", layout="wide")

# CSS
st.markdown("""<style>
.metric-card{background:linear-gradient(135deg,#1a1a2e,#16213e);border-radius:12px;padding:20px;text-align:center;border:1px solid #0f3460}
.metric-val{font-size:28px;font-weight:700;color:#e94560}
.metric-label{font-size:13px;color:#a0a0b0;margin-bottom:4px}
.status-badge{padding:4px 12px;border-radius:20px;font-size:12px;font-weight:600;display:inline-block}
.st-dtc{color:#22c55e}.st-lpakt{color:#f59e0b}.st-lkh{color:#3b82f6}.st-ht{color:#8b5cf6}
</style>""", unsafe_allow_html=True)

def fmt_money(v):
    try:
        v = int(float(v))
        if v >= 1e9: return f"{v/1e9:,.2f} tỷ"
        if v >= 1e6: return f"{v/1e6:,.1f} tr"
        return f"{v:,}"
    except: return "0"

def fmt_full(v):
    try: return f"{int(float(v)):,}"
    except: return "0"

def status_color(s):
    s = str(s).strip()
    colors = {'Đang thi công':'#22c55e','Lập PAKT-Tổng dự toán':'#f59e0b','Lập kế hoạch đầu thầu':'#3b82f6','Hoàn thành':'#8b5cf6'}
    return colors.get(s, '#94a3b8')

def parse_date_from_text(text, prefix):
    import re
    match = re.search(fr"{prefix}:\s*(\d{{1,2}})/(\d{{4}})", str(text), re.IGNORECASE)
    if match:
        return int(match.group(1)), int(match.group(2))
    return None, None

def analyze_project_health(row, current_year, current_month):
    status = str(row.get('Trạng thái', '')).strip()
    tien_do_text = str(row.get('Tiến độ', ''))
    khai_toan = float(row.get('Khái toán', 0)) if pd.notna(row.get('Khái toán')) else 0
    thuc_hien = float(row.get('Thực hiện', 0)) if pd.notna(row.get('Thực hiện')) else 0
    
    ty_le = (thuc_hien / khai_toan * 100) if khai_toan > 0 else 0
    
    if status in ['Hoàn thành', 'Nghiệm thu']:
        return "Hoàn thành", "🟢", f"Dự án đã hoàn thành. Giải ngân: {ty_le:.1f}%"
    
    kc_m, kc_y = parse_date_from_text(tien_do_text, 'KC')
    ht_m, ht_y = parse_date_from_text(tien_do_text, 'HT')
    
    if ht_m and ht_y:
        months_left = (ht_y - current_year) * 12 + (ht_m - current_month)
        if months_left < 0:
            return "Quá hạn", "🔴", f"Trễ hạn hoàn thành {-months_left} tháng (Hạn HT: {ht_m}/{ht_y}). Giải ngân mới đạt {ty_le:.1f}%"
        elif months_left <= 2 and ty_le < 30:
            return "Nguy cơ cao", "🟡", f"Chỉ còn {months_left} tháng đến hạn ({ht_m}/{ht_y}) nhưng giải ngân rất thấp ({ty_le:.1f}%)."
    
    if kc_m and kc_y:
        months_passed = (current_year - kc_y) * 12 + (current_month - kc_m)
        if months_passed > 2 and status in ['Lập PAKT-Tổng dự toán', 'Lập kế hoạch đầu thầu', 'Chưa xác định']:
            return "Trễ tiến độ", "🔴", f"Đã qua mốc khởi công ({kc_m}/{kc_y}) {months_passed} tháng nhưng vẫn ở trạng thái '{status}'."
        if months_passed > 1 and ty_le == 0 and status == 'Đang thi công':
            return "Cần lưu ý", "🟡", f"Khởi công từ {kc_m}/{kc_y} nhưng chưa có số liệu giải ngân (0%)."
            
    return "Đúng tiến độ", "🔵", f"Tiến độ bình thường. Giải ngân: {ty_le:.1f}%."

# Load data
try:
    import plotly.express as px
    import plotly.graph_objects as go
    HAS_PLOTLY = True
except ImportError:
    HAS_PLOTLY = False

df_th = load_tonghop()
pm_data = load_pm092()

# Merge PM_092 data
if not df_th.empty and 'Mã CT' in df_th.columns:
    df_th['Thực hiện PM'] = df_th['Mã CT'].map(lambda x: pm_data.get(str(x).strip(), 0))
    if 'Thực hiện' not in df_th.columns:
        df_th['Thực hiện'] = df_th['Thực hiện PM']
    else:
        df_th['Thực hiện'] = df_th.apply(lambda r: r['Thực hiện PM'] if r['Thực hiện PM'] > 0 else r['Thực hiện'], axis=1)

tab1, tab2 = st.tabs(["📊 Bảng số liệu chi tiết các dự án SCL", "📄 Bảng thuyết minh quyết toán"])

with tab1:
    col_export, _ = st.columns([3, 7])
    with col_export:
        if st.button("📊 Xuất báo cáo", type="primary", key="cloud_export"):
            try:
                from hangmuc_report import generate_hangmuc
                res = generate_hangmuc()
                if res is not None:
                    hangmuc_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'HangMuc.xlsx')
                    if os.path.exists(hangmuc_path):
                        with open(hangmuc_path, "rb") as f:
                            file_data = f.read()
                        st.download_button(
                            "📥 Tải HangMuc.xlsx", data=file_data,
                            file_name="HangMuc.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key="cloud_dl"
                        )
                else:
                    st.warning("Chưa có dữ liệu để xuất.")
            except Exception as e:
                st.error(f"Lỗi xuất báo cáo: {e}")

    if df_th.empty:
        st.warning("Chưa có dữ liệu. Vui lòng đặt file **Tổng hợp.xlsx** vào thư mục ứng dụng.")
    else:
        # Metrics
        total_ct = len(df_th)
        total_kh = int(df_th['Khái toán'].fillna(0).sum()) if 'Khái toán' in df_th.columns else 0
        total_th = int(df_th['Thực hiện'].fillna(0).sum()) if 'Thực hiện' in df_th.columns else 0
        total_qt = int(df_th['Quyết toán'].fillna(0).sum()) if 'Quyết toán' in df_th.columns else 0
        ty_le = (total_th/total_kh*100) if total_kh > 0 else 0

        m1,m2,m3,m4 = st.columns(4)
        m1.metric("Tổng Số Công Trình", f"{total_ct}")
        m2.metric("Tổng Giá Trị Khái Toán", f"{fmt_full(total_kh)} đ")
        m3.metric("Tổng Giá Trị Thực Hiện", f"{fmt_full(total_th)} đ")
        m4.metric("Tỷ Lệ Giải Ngân", f"{ty_le:.2f} %")

        # Risk Analysis
        now = datetime.datetime.now()
        current_y = now.year
        current_m = now.month
        
        health_data = []
        for idx, row in df_th.iterrows():
            ma = row.get('Mã CT', '')
            ten = row.get('Tên công trình', '')
            h_status, h_icon, h_insight = analyze_project_health(row, current_y, current_m)
            health_data.append({
                'Mã CT': ma,
                'Tên công trình': ten,
                'Trạng thái Sức khỏe': f"{h_icon} {h_status}",
                'Đánh giá & Khuyến nghị': h_insight,
                'Status_Raw': h_status
            })
            
        df_health = pd.DataFrame(health_data)
        
        st.markdown("#### 🚨 Đánh giá & Cảnh báo rủi ro (Executive Summary)")
        with st.container(border=True):
            r1, r2, r3, r4 = st.columns(4)
            r1.metric("🟢 Tốt / Hoàn thành", len(df_health[df_health['Status_Raw'] == 'Hoàn thành']))
            r2.metric("🔵 Đúng tiến độ", len(df_health[df_health['Status_Raw'] == 'Đúng tiến độ']))
            r3.metric("🟡 Nguy cơ cao / Lưu ý", len(df_health[df_health['Status_Raw'].isin(['Nguy cơ cao', 'Cần lưu ý'])]))
            r4.metric("🔴 Quá hạn / Trễ", len(df_health[df_health['Status_Raw'].isin(['Quá hạn', 'Trễ tiến độ'])]))
            
            st.dataframe(df_health.drop(columns=['Status_Raw']), width='stretch', hide_index=True)

        # Charts
        if HAS_PLOTLY:
            st.markdown("#### 📈 Sơ đồ trực quan hóa dữ liệu")
            ch1, ch2 = st.columns(2)

            with ch1:
                if 'Trạng thái' in df_th.columns:
                    status_counts = df_th['Trạng thái'].fillna('Chưa xác định').value_counts()
                    colors_map = {s: status_color(s) for s in status_counts.index}
                    fig1 = px.pie(values=status_counts.values, names=status_counts.index,
                        title="1. Tỷ trọng trạng thái dự án",
                        color=status_counts.index, color_discrete_map=colors_map, hole=0.4)
                    fig1.update_layout(
                        paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)',
                        font=dict(color='white',size=12),
                        legend=dict(orientation="v", yanchor="middle", y=0.5),
                        height=400, margin=dict(t=40,b=20,l=20,r=20))
                    fig1.update_traces(textposition='outside', textinfo='percent+label',
                        textfont_size=11, marker=dict(line=dict(color='#1a1a2e',width=2)))
                    st.plotly_chart(fig1, use_container_width=True, key='pie_chart')

            with ch2:
                if 'Khái toán' in df_th.columns:
                    top_df = df_th.nlargest(5, 'Khái toán').copy()
                    top_df['KT_ty'] = top_df['Khái toán'].fillna(0)/1e9
                    top_df['TH_ty'] = top_df['Thực hiện'].fillna(0)/1e9 if 'Thực hiện' in top_df.columns else 0
                    fig2 = go.Figure()
                    fig2.add_trace(go.Bar(name='Khái toán (Tỷ đ)', x=top_df['Mã CT'], y=top_df['KT_ty'],
                        marker_color='#3b82f6', text=[f"{v:.1f}" for v in top_df['KT_ty']], textposition='outside'))
                    fig2.add_trace(go.Bar(name='Thực hiện (Tỷ đ)', x=top_df['Mã CT'], y=top_df['TH_ty'],
                        marker_color='#f59e0b', text=[f"{v:.1f}" for v in top_df['TH_ty']], textposition='outside'))
                    fig2.update_layout(
                        title="2. Top dự án có mức ngân sách cao nhất", barmode='group',
                        paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)',
                        font=dict(color='white',size=12),
                        xaxis=dict(gridcolor='rgba(255,255,255,0.1)'),
                        yaxis=dict(gridcolor='rgba(255,255,255,0.1)', title='Tỷ đồng'),
                        height=400, margin=dict(t=40,b=20,l=20,r=20),
                        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1))
                    st.plotly_chart(fig2, use_container_width=True, key='bar_chart')

        # Table
        st.markdown("#### 📋 Bảng số liệu chi tiết các dự án SCL")
        display_cols = ['Mã CT','Tên công trình','Trạng thái','Khái toán','Thực hiện','Quyết toán']
        display_cols = [c for c in display_cols if c in df_th.columns]
        disp_df = df_th[display_cols].copy()
        for c in ['Khái toán','Thực hiện','Quyết toán']:
            if c in disp_df.columns:
                disp_df[c] = disp_df[c].fillna(0).astype(int)
        
        col_cfg = {}
        for c in ['Khái toán','Thực hiện','Quyết toán']:
            if c in disp_df.columns:
                col_cfg[c] = st.column_config.NumberColumn(format="%,d")
        
        st.dataframe(disp_df, width='stretch', hide_index=True, column_config=col_cfg)

        # Chi tiết từng công trình (chỉ xem)
        st.divider()
        st.markdown("#### 🔍 Xem chi tiết công trình")
        ct_names = []
        for _, r in df_th.iterrows():
            ma = str(r.get('Mã CT','')).strip()
            ten = str(r.get('Tên công trình','')).strip()
            ct_names.append(f"{ma} - {ten}")
        
        selected = st.selectbox("Chọn công trình:", ["-- Chọn để xem chi tiết --"] + ct_names)
        
        if selected != "-- Chọn để xem chi tiết --":
            sel_ma = selected.split(" - ")[0].strip()
            row_idx = df_th[df_th['Mã CT']==sel_ma].index
            if len(row_idx) > 0:
                row_data = df_th.loc[row_idx[0]]
                with st.container(border=True):
                    c1, c2, c3 = st.columns(3)
                    with c1:
                        st.write(f"**Tên công trình:** {row_data.get('Tên công trình','')}")
                        st.write(f"**Mã CT:** {sel_ma}")
                        st.write(f"**Trạng thái:** {row_data.get('Trạng thái','')}")
                    with c2:
                        st.write(f"**Khái toán:** {fmt_full(row_data.get('Khái toán',0))} đ")
                        st.write(f"**Thực hiện:** {fmt_full(row_data.get('Thực hiện',0))} đ")
                        st.write(f"**Quyết toán:** {fmt_full(row_data.get('Quyết toán',0))} đ")
                    with c3:
                        if pd.notna(row_data.get('Nội dung SCL')):
                            st.markdown(f"**Nội dung SCL:**\n\n{row_data.get('Nội dung SCL','')}")
                        if pd.notna(row_data.get('Tiến độ')):
                            st.markdown(f"**Tiến độ:**\n\n{row_data.get('Tiến độ','')}")

                # Hiển thị bảng quyết toán từ database nếu có
                db_df = load_db_data()
                if not db_df.empty:
                    ten_ct = str(row_data.get('Tên công trình', ''))
                    start_indices = db_df.index[db_df['Tên Công trình'] == ten_ct].tolist()
                    if not start_indices:
                        start_indices = db_df.index[db_df['Mã CT'].astype(str).str.strip() == sel_ma].tolist()
                    if start_indices:
                        start_idx = start_indices[0]
                        end_idx = len(db_df)
                        for i in range(start_idx + 1, len(db_df)):
                            val = str(db_df.at[i, 'STT']).strip().upper()
                            if val in ['I','II','III','IV','V','VI','VII','VIII','IX','X']:
                                end_idx = i
                                break
                        ct_data = db_df.iloc[start_idx:end_idx]
                        if len(ct_data) > 1:
                            st.markdown("**Bảng tổng hợp quyết toán:**")
                            sub_data = ct_data.iloc[1:][['STT','Tên Công trình','Giá trị Dự toán','Giá trị Q.định phê duyệt QT công trình']].copy()
                            sub_data = sub_data.rename(columns={'Tên Công trình':'Tên Hạng mục','Giá trị Q.định phê duyệt QT công trình':'Giá trị quyết toán'})
                            for c in ['Giá trị Dự toán','Giá trị quyết toán']:
                                sub_data[c] = pd.to_numeric(sub_data[c], errors='coerce').fillna(0).astype(int)
                            sub_data['Chênh lệch'] = sub_data['Giá trị Dự toán'] - sub_data['Giá trị quyết toán']
                            st.dataframe(sub_data, hide_index=True, width='stretch',
                                column_config={
                                    'Giá trị Dự toán': st.column_config.NumberColumn(format="%,d"),
                                    'Giá trị quyết toán': st.column_config.NumberColumn(format="%,d"),
                                    'Chênh lệch': st.column_config.NumberColumn(format="%,d"),
                                })

                # Hiển thị Chi tiết hồ sơ theo hạng mục
                st.divider()
                st.markdown("#### 📋 Chi tiết hồ sơ theo hạng mục")
                
                from form_module import load_chitiet_by_ma, load_hopdong_list
                
                t1, t2, t3, t4, t5, t6 = st.tabs([
                    "1. PAKT-Dự toán", "2. KH Đấu thầu", "3. KQ Đấu thầu", 
                    "4. Hợp đồng", "5. Vật tư", "6. Nghiệm thu - QT"
                ])
                
                with t1:
                    df_pakt = load_chitiet_by_ma('pakt_dt', sel_ma)
                    if not df_pakt.empty: st.dataframe(df_pakt, hide_index=True, width='stretch')
                    else: st.info("Chưa có dữ liệu.")
                
                with t2:
                    df_kh = load_chitiet_by_ma('kh_dau_thau', sel_ma)
                    if not df_kh.empty: st.dataframe(df_kh, hide_index=True, width='stretch')
                    else: st.info("Chưa có dữ liệu.")
                    
                with t3:
                    df_kq = load_chitiet_by_ma('kq_dau_thau', sel_ma)
                    if not df_kq.empty: st.dataframe(df_kq, hide_index=True, width='stretch')
                    else: st.info("Chưa có dữ liệu.")
                    
                with t4:
                    df_hd = load_hopdong_list(sel_ma)
                    if not df_hd.empty: st.dataframe(df_hd, hide_index=True, width='stretch')
                    else: st.info("Chưa có dữ liệu.")
                    
                with t5:
                    df_vt = load_chitiet_by_ma('vat_tu', sel_ma)
                    if not df_vt.empty: st.dataframe(df_vt, hide_index=True, width='stretch')
                    else: st.info("Chưa có dữ liệu.")
                    
                with t6:
                    df_nt = load_chitiet_by_ma('nghiem_thu_qt', sel_ma)
                    if not df_nt.empty: st.dataframe(df_nt, hide_index=True, width='stretch')
                    else: st.info("Chưa có dữ liệu.")

with tab2:
    st.header("📄 Bảng thuyết minh quyết toán")
    db_df_tab3 = load_db_data()
    
    if db_df_tab3.empty:
        st.info("Chưa có dữ liệu. Vui lòng cập nhật số liệu trước.")
    else:
        main_mask_tab3 = db_df_tab3['Kế hoạch'].notna()
        list_ct_tab3 = db_df_tab3.loc[main_mask_tab3, 'Tên Công trình'].dropna().unique().tolist()
        
        selected_ct_tab3 = st.selectbox("Chọn Công trình để xuất Thuyết minh QT:", ["-- Chọn --"] + list_ct_tab3, key="tmqt_select")
        
        if selected_ct_tab3 != "-- Chọn --":
            start_indices_t3 = db_df_tab3.index[db_df_tab3['Tên Công trình'] == selected_ct_tab3].tolist()
            if start_indices_t3:
                start_idx_t3 = start_indices_t3[0]
                end_idx_t3 = len(db_df_tab3)
                for i in range(start_idx_t3 + 1, len(db_df_tab3)):
                    val = str(db_df_tab3.at[i, 'STT']).strip().upper()
                    if val in ['I','II','III','IV','V','VI','VII','VIII','IX','X']:
                        end_idx_t3 = i
                        break
                
                ct_data_t3 = db_df_tab3.iloc[start_idx_t3:end_idx_t3]
                main_row_t3 = ct_data_t3.iloc[0]
                
                ten_ct_t3 = str(main_row_t3.get('Tên Công trình', ''))
                ma_ct_t3 = str(main_row_t3.get('Mã CT', '')) if pd.notna(main_row_t3.get('Mã CT')) else ''
                ke_hoach_t3 = main_row_t3.get('Kế hoạch', 0)
                if pd.isna(ke_hoach_t3): ke_hoach_t3 = 0
                ke_hoach_t3 = int(float(ke_hoach_t3))
                
                don_vi_ql_t3 = str(main_row_t3.get('Đơn vị QL', '')) if pd.notna(main_row_t3.get('Đơn vị QL')) else ''
                can_cu_pl_t3 = str(main_row_t3.get('Căn cứ pháp lý', '')) if pd.notna(main_row_t3.get('Căn cứ pháp lý')) else ''
                klcv_t3 = str(main_row_t3.get('Khối lượng công việc', '')) if pd.notna(main_row_t3.get('Khối lượng công việc')) else ''
                
                ngay_kc_t3 = main_row_t3.get('Ngày khởi công')
                ngay_ht_t3 = main_row_t3.get('Ngày hoàn thành')
                
                def format_date_vn(d):
                    if pd.isna(d) or d is None: return '....../....../...........'
                    if isinstance(d, pd.Timestamp): d = d.date()
                    if isinstance(d, (datetime.date, datetime.datetime)): return d.strftime('%d/%m/%Y')
                    return str(d)
                
                ngay_kc_str = format_date_vn(ngay_kc_t3)
                ngay_ht_str = format_date_vn(ngay_ht_t3)
                
                gt_dt_scl = 0
                gt_qt_scl = 0
                for idx_t3, row_t3 in ct_data_t3.iterrows():
                    stt_val = str(row_t3['STT']).strip().upper()
                    if stt_val == 'SCL':
                        dt_val = row_t3.get('Giá trị Dự toán', 0)
                        qt_val = row_t3.get('Giá trị Q.định phê duyệt QT công trình', 0)
                        if pd.notna(dt_val): gt_dt_scl = int(float(dt_val))
                        if pd.notna(qt_val): gt_qt_scl = int(float(qt_val))
                        break
                
                st.subheader("Xem trước thông tin")
                with st.container(border=True):
                    c1, c2 = st.columns(2)
                    with c1:
                        st.write(f"**Tên công trình:** {ten_ct_t3}")
                        st.write(f"**Mã CT:** {ma_ct_t3}")
                        st.write(f"**Giá trị kế hoạch vốn:** {f'{ke_hoach_t3:,}'} đồng")
                        st.write(f"**Giá trị dự toán được duyệt:** {f'{gt_dt_scl:,}'} đồng")
                    with c2:
                        st.write(f"**Thời gian khởi công:** {ngay_kc_str}")
                        st.write(f"**Thời gian hoàn thành:** {ngay_ht_str}")
                        st.write(f"**Giá trị quyết toán:** {f'{gt_qt_scl:,}'} đồng")
                        st.write(f"**Đơn vị QL:** {don_vi_ql_t3}")
                
                if can_cu_pl_t3:
                    st.write("**Căn cứ pháp lý:**")
                    st.info(can_cu_pl_t3)
                if klcv_t3:
                    st.write("**Khối lượng công việc:**")
                    st.info(klcv_t3)
                
                st.divider()
                
                if st.button("📥 Xuất file Word - Bảng thuyết minh quyết toán", type="primary", key="btn_tmqt"):
                    try:
                        from docx import Document as DocxDocument
                        from docx.shared import Pt, Cm
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        
                        doc = DocxDocument('Mẫu TMQT.docx')
                        ghi_chu_t3 = str(main_row_t3.get('Ghi chú', '')) if pd.notna(main_row_t3.get('Ghi chú')) else ''
                        now = datetime.datetime.now()
                        
                        def format_money(val):
                            if val == 0: return '0'
                            return f'{val:,}'.replace(',', '.')
                        
                        def replace_para_with_lines(p, lines):
                            if not lines:
                                p.text = ""
                                return
                            style = p.style
                            left_indent = p.paragraph_format.left_indent
                            first_line_indent = p.paragraph_format.first_line_indent
                            for line in lines[:-1]:
                                new_p = p.insert_paragraph_before(line, style=style)
                                new_p.paragraph_format.left_indent = left_indent
                                new_p.paragraph_format.first_line_indent = first_line_indent
                                new_p.paragraph_format.space_after = Pt(0)
                                new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            p.text = lines[-1]
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        paragraphs = list(doc.paragraphs)
                        for p in paragraphs:
                            p.paragraph_format.space_after = Pt(0)
                            text_val = p.text.strip()
                            if not text_val: continue
                            
                            if "- Tên danh mục:" in text_val:
                                lines = [f"- Tên danh mục: {ten_ct_t3}", f"- Mã công trình: {ma_ct_t3}"]
                                replace_para_with_lines(p, lines)
                            elif "- Giá trị vốn kế hoạch:" in text_val:
                                p.text = f"- Giá trị vốn kế hoạch: {format_money(ke_hoach_t3)} đồng"
                            elif "sửa chữa lớn năm" in text_val:
                                p.text = f"- Thuộc kế hoạch vốn sửa chữa lớn năm {now.year}"
                            elif "Hình thức tự làm hay thuê ngoài" in text_val:
                                p.text = f"- Hình thức tự làm hay thuê ngoài: {ghi_chu_t3}"
                            elif "- Tên đơn vị thi công" in text_val:
                                p.text = f"- Tên đơn vị thi công: {don_vi_ql_t3}"
                            elif "- Giá trị dự toán được duyệt" in text_val:
                                p.text = f"- Giá trị dự toán được duyệt: {format_money(gt_dt_scl)} đồng"
                            elif "- Thời gian khởi công" in text_val:
                                p.text = f"- Thời gian khởi công: {ngay_kc_str}"
                            elif "- Thời gian hoàn thành" in text_val:
                                p.text = f"- Thời gian hoàn thành: {ngay_ht_str}"
                            elif "- Giá trị quyết toán" in text_val and "hoàn thành" in text_val:
                                p.text = f"- Giá trị quyết toán danh mục hoàn thành: {format_money(gt_qt_scl)} đồng"
                            elif "Khối lượng công việc chủ yếu đã tiến hành" in text_val:
                                lines = [f"- Khối lượng công việc chủ yếu đã tiến hành (thay thế, sửa chữa những bộ phận nào của TSCĐ):"]
                                if klcv_t3:
                                    lines.extend([line for line in klcv_t3.split('\n') if line.strip()])
                                replace_para_with_lines(p, lines)
                            elif "Các căn cứ về chế độ để lập quyết toán" in text_val:
                                lines = [f"- Các căn cứ về chế độ để lập quyết toán:"]
                                if can_cu_pl_t3:
                                    lines.extend([line for line in can_cu_pl_t3.split('\n') if line.strip()])
                                replace_para_with_lines(p, lines)
                            elif "+ .........." in text_val:
                                p.text = ""
                            elif "ngày       tháng      năm" in text_val:
                                p.text = text_val.replace("2026", str(now.year))
                        
                        for section in doc.sections:
                            section.top_margin = Cm(2)
                            section.bottom_margin = Cm(2)
                            section.left_margin = Cm(3)
                            section.right_margin = Cm(2)
                        for p in doc.paragraphs:
                            for run in p.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                        for t in doc.tables:
                            for row in t.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        p.paragraph_format.space_after = Pt(0)
                                        for run in p.runs:
                                            run.font.name = 'Times New Roman'
                                            run.font.size = Pt(12)
                        
                        output_docx = BytesIO()
                        doc.save(output_docx)
                        docx_data = output_docx.getvalue()
                        
                        safe_name = ten_ct_t3[:30].replace('/', '_').replace('\\', '_').replace(':', '_')
                        st.download_button(
                            label="📥 Tải xuống file Word",
                            data=docx_data,
                            file_name=f"Thuyet_minh_QT_{safe_name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_tmqt_final"
                        )
                        st.success("✅ Đã tạo file Word thành công!")
                    except Exception as e:
                        st.error(f"Lỗi khi tạo file Word: {e}")
